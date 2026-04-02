"""Parse documents (TXT, EPUB, PDF, DOCX) into chapters."""
import re
from pathlib import Path
from dataclasses import dataclass


@dataclass
class ParsedChapter:
    title: str
    text: str


# Common chapter heading patterns
CHAPTER_PATTERNS = [
    re.compile(r"^(chapter\s+\d+[.:;]?\s*.*)$", re.IGNORECASE | re.MULTILINE),
    re.compile(r"^(CHAPTER\s+[IVXLCDM]+[.:;]?\s*.*)$", re.MULTILINE),
    re.compile(r"^(Part\s+\d+[.:;]?\s*.*)$", re.IGNORECASE | re.MULTILINE),
    re.compile(r"^(PART\s+[IVXLCDM]+[.:;]?\s*.*)$", re.MULTILINE),
    re.compile(r"^(Section\s+\d+[.:;]?\s*.*)$", re.IGNORECASE | re.MULTILINE),
    re.compile(r"^(\d+\.\s+[A-Z].{2,60})$", re.MULTILINE),
]


def parse_text(text: str) -> list[ParsedChapter]:
    """Split plain text into chapters by detecting headings."""
    text = text.strip()
    if not text:
        return []

    # Try each chapter pattern
    for pattern in CHAPTER_PATTERNS:
        matches = list(pattern.finditer(text))
        if len(matches) >= 2:
            return _split_by_matches(text, matches)

    # Fallback: split by triple newlines or large gaps
    parts = re.split(r"\n{3,}", text)
    if len(parts) >= 2:
        chapters = []
        for i, part in enumerate(parts):
            part = part.strip()
            if not part:
                continue
            # Use first line as title if short enough
            lines = part.split("\n", 1)
            if len(lines[0]) < 80 and len(lines) > 1:
                title = lines[0].strip()
                body = lines[1].strip()
            else:
                title = f"Section {i + 1}"
                body = part
            chapters.append(ParsedChapter(title=title, text=body))
        return chapters if len(chapters) >= 2 else [ParsedChapter(title="Full Text", text=text)]

    # Single chapter
    return [ParsedChapter(title="Full Text", text=text)]


def _split_by_matches(text: str, matches: list) -> list[ParsedChapter]:
    """Split text using regex match positions."""
    chapters = []
    for i, match in enumerate(matches):
        title = match.group(1).strip()
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        if body:
            chapters.append(ParsedChapter(title=title, text=body))
    return chapters


def parse_epub(file_path: Path) -> list[ParsedChapter]:
    """Parse EPUB file into chapters."""
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    book = epub.read_epub(str(file_path))
    chapters = []

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        text = soup.get_text(separator="\n").strip()
        if not text or len(text) < 50:
            continue

        # Try to get title from headings
        heading = soup.find(["h1", "h2", "h3"])
        if heading:
            title = heading.get_text().strip()
        else:
            # Use first line
            first_line = text.split("\n")[0].strip()
            title = first_line[:80] if first_line else f"Chapter {len(chapters) + 1}"

        chapters.append(ParsedChapter(title=title, text=text))

    if not chapters:
        # Fallback: extract all text and use text parser
        all_text = ""
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "html.parser")
            all_text += soup.get_text(separator="\n") + "\n\n"
        return parse_text(all_text)

    return chapters


def parse_pdf(file_path: Path) -> list[ParsedChapter]:
    """Parse PDF file into chapters."""
    import fitz  # pymupdf

    doc = fitz.open(str(file_path))
    full_text = ""
    for page in doc:
        full_text += page.get_text() + "\n"
    doc.close()

    return parse_text(full_text)


def parse_docx(file_path: Path) -> list[ParsedChapter]:
    """Parse DOCX file into chapters using heading styles."""
    try:
        from docx import Document
    except ImportError:
        # Fallback: try to extract text via pymupdf or basic XML parsing
        return _parse_docx_basic(file_path)

    doc = Document(str(file_path))
    chapters: list[ParsedChapter] = []
    current_title = ""
    current_text = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            current_text += "\n"
            continue

        # Detect heading styles
        is_heading = para.style and para.style.name and "Heading" in para.style.name

        if is_heading:
            # Save previous chapter
            if current_text.strip():
                title = current_title or f"Chapter {len(chapters) + 1}"
                chapters.append(ParsedChapter(title=title, text=current_text.strip()))
            current_title = text
            current_text = ""
        else:
            current_text += text + "\n"

    # Save last chapter
    if current_text.strip():
        title = current_title or f"Chapter {len(chapters) + 1}"
        chapters.append(ParsedChapter(title=title, text=current_text.strip()))

    if not chapters:
        # No headings found — use text parser
        all_text = "\n".join(p.text for p in doc.paragraphs)
        return parse_text(all_text)

    return chapters


def _parse_docx_basic(file_path: Path) -> list[ParsedChapter]:
    """Basic DOCX parsing without python-docx by extracting XML text."""
    import zipfile
    from xml.etree import ElementTree

    with zipfile.ZipFile(str(file_path)) as z:
        with z.open("word/document.xml") as f:
            tree = ElementTree.parse(f)

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = []
    for p in tree.iter(f"{{{ns['w']}}}p"):
        texts = [t.text for t in p.iter(f"{{{ns['w']}}}t") if t.text]
        if texts:
            paragraphs.append("".join(texts))

    return parse_text("\n\n".join(paragraphs))


def parse_file(file_path: Path) -> list[ParsedChapter]:
    """Auto-detect file type and parse into chapters."""
    ext = file_path.suffix.lower()
    if ext == ".epub":
        return parse_epub(file_path)
    elif ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext in (".txt", ".text", ".md"):
        return parse_text(file_path.read_text(encoding="utf-8", errors="replace"))
    else:
        # Try as plain text
        return parse_text(file_path.read_text(encoding="utf-8", errors="replace"))
